#!/usr/bin/env python3
"""Cross-check manuscript tables + figures against the latest run JSON."""
import hashlib
import json
import os
from pathlib import Path

import docx

d = json.load(open("results/analysis_results.json"))
m = d["main"]
tiers = m["tiers"]
sys_ = m["system"]
gc = d["grace_comparison"]
cal = m["calibration"]
miss = d["missingness"]
trade = {r["threshold"]: r for r in d["threshold_tradeoff"]}
doc = docx.Document("manuscript/acs_mortality_triage_manuscript.docx")

print("=== TABEL DI DOCX ===")
for ti, t in enumerate(doc.tables, 1):
    print(f"\n--- Table {ti} ({len(t.rows)} baris x {len(t.columns)} kolom) ---")
    for r in t.rows[:4]:
        print("  |", " | ".join(c.text.strip()[:22] for c in r.cells))

print("\n=== CEK ANGKA TABEL vs JSON ===")
checks = [
    # (deskripsi, harus ada di docx text semua tabel, nilai dari JSON)
    ("Tiers HIGH n", str(tiers["high"]["n"]), f"{tiers['high']['deaths']}"),
    ("Tiers INT n", str(tiers["intermediate"]["n"]), str(tiers["intermediate"]["deaths"])),
    ("Tiers LOW n", str(tiers["low"]["n"]), str(tiers["low"]["deaths"])),
    ("Not flagged n", str(tiers["not_flagged"]["n"]), str(tiers["not_flagged"]["deaths"])),
    ("CM TP/FP", str(sys_["tp"]), str(sys_["fp"])),
    ("CM FN/TN", str(sys_["fn"]), str(sys_["tn"])),
    ("AUC model", f"{m['auc']:.3f}", ""),
    ("AUC GRACE", f"{gc['grace_auc']:.3f}", ""),
    ("cal slope", f"{cal['slope']:.3f}", ""),
    ("cal ece", f"{cal['ece']:.3f}", ""),
    ("miss complete", f"{miss['complete_case_n']:,}", ""),
    ("miss any", str(miss["any_missing_n"]), ""),
]
full = "\n".join(p.text for p in doc.paragraphs)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            full += "\n" + c.text
for label, v1, v2 in checks:
    ok1 = v1 in full
    ok2 = (not v2) or (v2 in full)
    print(f"  {'OK ' if ok1 and ok2 else 'FAIL'} | {label}: {v1}{'/' + v2 if v2 else ''}")

print("\n=== TRADE-OFF TABEL S2 ===")
for th, row in trade.items():
    print(f"  {th}: sens {row['system']['sensitivity']*100:.1f}% FP {row['false_positives']} "
          f"flagged {row['flagged_n']} escal {row['escalated_n']} PPV {row['system']['ppv']*100:.1f}%")

print("\n=== GAMBAR: FILE vs EMBED DOCX ===")
def sha(p):
    return hashlib.md5(open(p, "rb").read()).hexdigest()

fig_dir = Path("figures")
media = Path("manuscript")  # docx is zip; extract via docx internals
# python-docx: embedded images live in package parts
img_parts = {}
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        img_parts[rel.target_ref.split("/")[-1]] = rel.target_part.blob
print(f"  gambar ter-embed di docx: {len(img_parts)}")
for fig in sorted(fig_dir.glob("*.png")):
    fh = sha(fig)
    matched = [name for name, blob in img_parts.items() if hashlib.md5(blob).hexdigest() == fh]
    print(f"  {'OK ' if matched else 'BEDA'} | {fig.name} ({fig.stat().st_size//1024}KB) mtime={fig.stat().st_mtime:.0f}"
          f" {'-> embed: ' + matched[0] if matched else '(TIDAK cocok dgn embed docx!)'}")

print("\n  mtime JSON analysis:", os.path.getmtime("results/analysis_results.json"))
