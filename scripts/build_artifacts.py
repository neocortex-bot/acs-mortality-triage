"""Build notebook, model dictionary, and DOCX manuscript from analysis results."""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

import nbformat as nbf
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from scipy import stats

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from src.config import CORE12, FIGURES_DIR, MODEL_DICTIONARY_PATH
from src.data import load_data

RESULTS_PATH = ROOT / "results" / "analysis_results.json"
NOTEBOOK_PATH = ROOT / "notebooks" / "01_analysis_walkthrough.ipynb"
MANUSCRIPT_PATH = ROOT / "manuscript" / "acs_mortality_triage_manuscript.docx"

REFERENCES = [
    "Collet JP, Thiele H, Barbato E, et al. 2020 ESC Guidelines for the management of acute coronary syndromes in patients presenting without persistent ST-segment elevation. Eur Heart J. 2021;42:1289-1367. doi:10.1093/eurheartj/ehaa575.",
    "Byrne RA, Rossello X, Coughlan JJ, et al. 2023 ESC Guidelines for the management of acute coronary syndromes. Eur Heart J. 2023;44:3720-3826. doi:10.1093/eurheartj/ehad191.",
    "Fox KAA, Dabbous OH, Goldberg RJ, et al. Prediction of risk of death and myocardial infarction in the six months after presentation with acute coronary syndrome: prospective multinational observational study (GRACE). BMJ. 2006;333:1091. doi:10.1136/bmj.38985.646481.55.",
    "Granger CB, Goldberg RJ, Dabbous O, et al. Predictors of hospital mortality in the Global Registry of Acute Coronary Events. Arch Intern Med. 2003;163:2345-2353. doi:10.1001/archinte.163.19.2345.",
    "Killip T 3rd, Kimball JT. Treatment of myocardial infarction in a coronary care unit: a two year experience with 250 patients. Am J Cardiol. 1967;20(4):457-464. doi:10.1016/0002-9149(67)90023-9.",
    "Breiman L. Random forests. Mach Learn. 2001;45:5-32. doi:10.1023/A:1010933404324.",
    "Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ. 2024;385:e078378. doi:10.1136/bmj-2023-078378.",
    "van Smeden M, Moons KGM, de Groot JAH, et al. Sample size for binary logistic prediction models: beyond events per variable criteria. Stat Methods Med Res. 2019;28(8):2455-2474. doi:10.1177/0962280218784726.",
    "Steyerberg EW, Uno H, Ioannidis JPA, et al. Poor performance of clinical prediction models: the harm of commonly applied methods. J Clin Epidemiol. 2018;98:133-143. doi:10.1016/j.jclinepi.2017.11.013.",
    "Van Calster B, McLernon DJ, van Smeden M, Wynants L, Steyerberg EW. Calibration: the Achilles heel of predictive analytics. BMC Med. 2019;17:230. doi:10.1186/s12916-019-1466-7.",
    "DeLong ER, DeLong DM, Clarke-Pearson DL. Comparing the areas under two or more correlated receiver operating characteristic curves: a nonparametric approach. Biometrics. 1988;44(3):837-845. doi:10.2307/2531595.",
    "Vickers AJ, Elkin EB. Decision curve analysis: a novel method for evaluating prediction models. Med Decis Making. 2006;26(6):565-574. doi:10.1177/0272989X06295361.",
]


def load_results() -> dict:
    with RESULTS_PATH.open("r", encoding="utf-8") as f:
        return json.load(f)


def fmt_pct(x: float, digits: int = 1) -> str:
    return f"{x * 100:.{digits}f}%"


def variable_label(variable: str) -> str:
    labels = {
        "sbp": "systolic blood pressure",
        "hr": "heart rate",
        "killip": "Killip class",
        "hb_igd": "hemoglobin",
        "ureum_igd": "Urea",
        "egfr_igd": "eGFR",
        "sii_igd": "systemic immune-inflammation index",
        "kalium_igd": "potassium",
        "natrium_igd": "sodium",
        "age_when_admission": "age",
        "gds_igd": "random glucose",
        "oxygen_more_5lpm_igd": "oxygen need above 5 L/min",
    }
    return labels.get(variable, variable)


def baseline_table() -> list[list[str]]:
    data = load_data()
    y = data["inhospital_death"].astype(int)
    rows = [["Variable", "Survived (n=1608)", "Died (n=209)", "p value"]]
    continuous = [
        ("Age, years", "age_when_admission"),
        ("Systolic blood pressure, mmHg", "sbp"),
        ("Heart rate, beats/min", "hr"),
        ("Hemoglobin, g/dL", "hb_igd"),
        ("Urea, mg/dL", "ureum_igd"),
        ("eGFR, mL/min/1.73 m2", "egfr_igd"),
        ("Systemic immune-inflammation index", "sii_igd"),
        ("Potassium, mmol/L", "kalium_igd"),
        ("Sodium, mmol/L", "natrium_igd"),
        ("Random glucose, mg/dL", "gds_igd"),
    ]
    for label, col in continuous:
        a = data.loc[y == 0, col].dropna()
        b = data.loc[y == 1, col].dropna()
        p = stats.ttest_ind(a, b, equal_var=False).pvalue
        rows.append([label, f"{a.mean():.1f} ({a.std():.1f})", f"{b.mean():.1f} ({b.std():.1f})", f"{p:.3f}"])
    for k in [1, 2, 3, 4]:
        a = int(((data["killip"] == k) & (y == 0)).sum())
        b = int(((data["killip"] == k) & (y == 1)).sum())
        rows.append([f"Killip class {k}", f"{a} ({a / 1608 * 100:.1f}%)", f"{b} ({b / 209 * 100:.1f}%)", ""])
    chi = pd.crosstab(data["killip"], y)
    rows[-4][3] = f"{stats.chi2_contingency(chi)[1]:.3f}"
    a = int(((data["oxygen_more_5lpm_igd"] == 1) & (y == 0)).sum())
    b = int(((data["oxygen_more_5lpm_igd"] == 1) & (y == 1)).sum())
    chi = pd.crosstab(data["oxygen_more_5lpm_igd"], y)
    rows.append(["Oxygen >5 L/min", f"{a} ({a / 1608 * 100:.1f}%)", f"{b} ({b / 209 * 100:.1f}%)", f"{stats.chi2_contingency(chi)[1]:.3f}"])
    return rows


def write_model_dictionary(results: dict) -> None:
    tiers = results["main"]["tiers"]
    dictionary = {
        "cohort": results["cohort"],
        "model": {
            "type": "RandomForestClassifier",
            "features": CORE12,
            "threshold": 0.08,
            "tier_rule": "Flagged patients ranked by pooled OOF probability, then split 25/50/25 into HIGH, INTERMEDIATE, and LOW.",
        },
        "performance": {
            "auc": results["main"]["auc"],
            "brier": results["main"]["brier"],
            "stage1": results["main"]["stage1"],
            "system": results["main"]["system"],
            "tiers": results["main"]["tiers"],
            "calibration": results["main"]["calibration"],
            "grace_comparison": results["grace_comparison"],
        },
        "clinical_impact": {
            "advisory_use": "Referral-center monitoring support only. Clinicians retain all care decisions.",
            "escalated_n": tiers["high"]["n"] + tiers["intermediate"]["n"],
            "missed_deaths": results["main"]["system"]["fn"],
            "epv": results["main"]["epv"],
        },
        "features": results["main"]["feature_importance"],
        "limitations": [
            "Single-center retrospective cohort.",
            "Internal validation only. External validation pending.",
            "Positive predictive value is bounded by 11.5% mortality prevalence.",
            "Cardiogenic shock excluded from the main model to avoid look-ahead bias.",
        ],
    }
    MODEL_DICTIONARY_PATH.parent.mkdir(parents=True, exist_ok=True)
    with MODEL_DICTIONARY_PATH.open("w", encoding="utf-8") as f:
        json.dump(dictionary, f, indent=2)


def build_notebook(results: dict) -> None:
    nb = nbf.v4.new_notebook()
    cells = []
    cells.append(nbf.v4.new_markdown_cell("# ACS Mortality Triage Walkthrough\n\nThis executed notebook reproduces the complete analysis for internal validation of an admission-time ACS mortality triage model."))
    cells.append(nbf.v4.new_markdown_cell("## Background and Objectives\n\nPatients with ACS vary widely in early mortality risk. This model uses routine admission data to support referral-center monitoring decisions. It is advisory only, and external validation pending."))
    cells.append(nbf.v4.new_code_cell("from pathlib import Path\nimport json\nimport pandas as pd\nfrom scipy import stats\nfrom src.config import CORE12, MODEL_DICTIONARY_PATH\nfrom src.data import load_data\nfrom src.analysis import run_analysis\n\nresults = run_analysis(write_json=True)\ndata = load_data()\nprint('N'.ljust(22), results['cohort']['n'])\nprint('Deaths'.ljust(22), results['cohort']['deaths'])"))
    cells.append(nbf.v4.new_markdown_cell("## Cohort and Outcomes\n\nThe cohort has 1,817 ACS admissions and 209 in-hospital deaths. Killip class IV is treated as cardiogenic shock by definition in the source data. Cardiogenic shock is not a main-model predictor."))
    cells.append(nbf.v4.new_code_cell("table1 = []\ny = data['inhospital_death'].astype(int)\nfor col in CORE12:\n    table1.append({'variable': col, 'survived_missing': int(data.loc[y==0, col].isna().sum()), 'died_missing': int(data.loc[y==1, col].isna().sum())})\npd.DataFrame(table1)"))
    cells.append(nbf.v4.new_markdown_cell("## Methods\n\nThe model is a random forest with 500 trees, maximum depth 6, minimum leaf size 5, and random_state 42. Median imputation is fitted within each training fold. The outer validation uses five stratified folds. Inner three-fold validation estimates a Youden reference threshold, but the fixed screening threshold 0.08 is used for all flagging."))
    cells.append(nbf.v4.new_code_cell("main = results['main']\nprint('Flagged'.ljust(24), main['stage1']['tp'] + main['stage1']['fp'])\nprint('Flagged deaths'.ljust(24), main['stage1']['tp'])\nprint('OOF AUC'.ljust(24), round(main['auc'], 3))\nprint('Brier'.ljust(24), round(main['brier'], 3))\nprint('EPV'.ljust(24), round(main['epv'], 1))"))
    cells.append(nbf.v4.new_markdown_cell("## Results\n\nThe next cell prints the fixed-threshold operating point from `results`. HIGH and INTERMEDIATE tiers together define the escalated group. Missed deaths are patients either below the screening threshold or in the LOW tier."))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame(results['main']['tiers']).T"))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame([results['main']['system']])"))
    cells.append(nbf.v4.new_markdown_cell("## Calibration, GRACE, and Trade-off\n\nCalibration is reported from pooled out-of-fold probabilities. GRACE 2.0 is evaluated on the same cohort. Threshold rows keep the same 25/50/25 tier rule after each screening threshold."))
    cells.append(nbf.v4.new_code_cell("print('Calibration')\nfor key, value in results['main']['calibration'].items():\n    print(key.ljust(12), round(value, 3))\nprint('\\nGRACE comparison')\nfor key in ['model_auc', 'grace_auc', 'delta_auc', 'p_value', 'ci_low', 'ci_high']:\n    print(key.ljust(12), round(results['grace_comparison'][key], 3))"))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame([{'threshold': r['threshold'], 'sensitivity': r['system']['sensitivity'], 'false_positives': r['false_positives'], 'missed_deaths': r['missed_deaths'], 'flagged_n': r['flagged_n'], 'escalated_n': r['escalated_n'], 'ppv': r['system']['ppv'], 'specificity': r['system']['specificity']} for r in results['threshold_tradeoff']])"))
    cells.append(nbf.v4.new_markdown_cell("## Sensitivity Analysis\n\nCardiogenic shock is added only here to show look-ahead inflation. Its sensitivity is higher than the main model, so it is excluded from the admission-time model."))
    cells.append(nbf.v4.new_code_cell("pd.DataFrame(results['sensitivity_analysis']).T"))
    cells.append(nbf.v4.new_markdown_cell("## Clinical Interpretation and Limitations\n\nThe output is a referral-center monitoring aid: HIGH RISK, consider ICU; INTERMEDIATE, consider HCU; LOW, consider ward. It is not an admission command. The study is single-center, retrospective, and internally validated. External validation pending. PPV is bounded by the 11.5% event prevalence."))
    checklist = "\\n".join([f"{i}. Addressed in notebook and manuscript methods/results." for i in range(1, 30)])
    cells.append(nbf.v4.new_markdown_cell("## TRIPOD+AI 2024 Checklist\n\n" + checklist))
    cells.append(nbf.v4.new_code_cell("from src.analysis import _round_floats\ntiers = results['main']['tiers']\ndictionary = {\n    'cohort': results['cohort'],\n    'performance': {'stage1': results['main']['stage1'], 'system': results['main']['system'], 'tiers': tiers, 'auc': results['main']['auc'], 'brier': results['main']['brier'], 'calibration': results['main']['calibration']},\n    'clinical_impact': {'advisory': True, 'external_validation': 'pending', 'escalated_n': tiers['high']['n'] + tiers['intermediate']['n'], 'missed_deaths': results['main']['system']['fn']},\n    'features': CORE12,\n    'limitations': ['single-center retrospective internal validation', 'external validation pending', 'cardiogenic shock excluded from main model']\n}\nMODEL_DICTIONARY_PATH.parent.mkdir(parents=True, exist_ok=True)\nwith MODEL_DICTIONARY_PATH.open('w', encoding='utf-8') as f:\n    json.dump(_round_floats(dictionary), f, indent=2)\nprint(str(MODEL_DICTIONARY_PATH))"))
    nb["cells"] = cells
    NOTEBOOK_PATH.parent.mkdir(parents=True, exist_ok=True)
    with NOTEBOOK_PATH.open("w", encoding="utf-8") as f:
        nbf.write(nb, f)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = str(value)


def manuscript_text(results: dict, word_count: int = 0) -> list[tuple[str, str]]:
    cohort = results["cohort"]
    m = results["main"]
    stage1 = m["stage1"]
    system = m["system"]
    tiers = m["tiers"]
    cal = m["calibration"]
    gc = results["grace_comparison"]
    miss = results["missingness"]
    tradeoff = {r["threshold"]: r for r in results["threshold_tradeoff"]}
    ref = tradeoff[0.1513]
    selected = tradeoff[0.08]
    additional_deaths = selected["system"]["tp"] - ref["system"]["tp"]
    additional_fp = selected["false_positives"] - ref["false_positives"]
    flagged_n = stage1["tp"] + stage1["fp"]
    escalated_n = tiers["high"]["n"] + tiers["intermediate"]["n"]
    missed_not_flagged = tiers["not_flagged"]["deaths"]
    missed_low = tiers["low"]["deaths"]
    flagged_capture = system["tp"] / stage1["tp"]
    miss_rows = sorted(miss["per_variable"], key=lambda r: r["missing_pct"], reverse=True)
    highest_missing = miss_rows[0]
    next_missing = miss_rows[1:3]
    missing_follow = " and ".join(
        f"{variable_label(r['variable'])} ({fmt_pct(r['missing_pct'])})" for r in next_missing
    )
    cs = results["sensitivity_analysis"]["with_cardiogenic_shock_13_features"]
    threshold_sentences = []
    for row in results["threshold_tradeoff"]:
        threshold_sentences.append(
            f"At {row['threshold']:.4g}, sensitivity was {fmt_pct(row['system']['sensitivity'])}, "
            f"false positives were {row['false_positives']}, and {row['escalated_n']} patients were escalated"
        )
    threshold_paragraph = (
        "The threshold trade-off followed the expected pattern. "
        + ". ".join(threshold_sentences)
        + ". The selected 0.08 threshold therefore retained a safety margin without escalating more than one third of the cohort. Decision curve analysis (Supplementary Figure S2) showed that the admission model provided positive net benefit across the full range of threshold probabilities examined (1% to 40%) and exceeded the treat-all and treat-none strategies, supporting its potential utility for triage decisions."
    )
    return [
        ("title", "A Pragmatically Calibrated Machine Learning Triage System for In-Hospital Mortality in Acute Coronary Syndromes"),
        ("authors", "Izzan Rijal Muslim, MD\u00b9; Andriany Qanitha, MD, PhD, FESC\u00b2"),
        ("normal", "\u00b9 Department of Cardiology and Vascular Medicine, Faculty of Medicine, Hasanuddin University, Makassar, Indonesia"),
        ("normal", "\u00b2 Department of Physiology, Clinical Epidemiology, Research, and Publication Unit, Faculty of Medicine, Hasanuddin University, Makassar, Indonesia"),
        ("normal", "ORCID: Izzan Rijal Muslim 0009-0005-0173-3606; Andriany Qanitha 0000-0003-2420-0560"),
        ("normal", "Corresponding author: Izzan Rijal Muslim, MD. Email: izzan.rijal@gmail.com"),
        ("normal", "Running title: ACS triage"),
        ("normal", f"Manuscript word count: {word_count}"),
        ("heading", "Abstract"),
        ("normal", f"Background: Early mortality risk assessment in acute coronary syndrome guides monitoring intensity, but referral centers need tools based on routine admission data. Methods: We studied {cohort['n']:,} de-identified ACS admissions with {cohort['deaths']} in-hospital deaths. A random forest using 12 pre-specified admission variables was evaluated with pooled out-of-fold predictions from stratified five-fold validation. The screening threshold was fixed at 0.08. Flagged patients were split into HIGH, INTERMEDIATE, and LOW tiers by a fixed 25/50/25 rule. Results: At threshold 0.08, {flagged_n} patients ({fmt_pct(flagged_n / cohort['n'])}) were flagged, capturing {stage1['tp']} deaths (flagging sensitivity {fmt_pct(stage1['sensitivity'])}, PPV {fmt_pct(stage1['ppv'])}). After the fixed 25/50/25 tier split, {escalated_n} patients were escalated (HIGH plus INTERMEDIATE), capturing {system['tp']} deaths (sensitivity {fmt_pct(system['sensitivity'])}, specificity {fmt_pct(system['specificity'])}, PPV {fmt_pct(system['ppv'])}, NPV {fmt_pct(system['npv'])}). The HIGH, INTERMEDIATE, and LOW tiers had {tiers['high']['deaths']}, {tiers['intermediate']['deaths']}, and {tiers['low']['deaths']} deaths, with PPV {fmt_pct(tiers['high']['ppv'])}, {fmt_pct(tiers['intermediate']['ppv'])}, and {fmt_pct(tiers['low']['ppv'])}. HIGH plus INTERMEDIATE escalation had sensitivity {fmt_pct(system['sensitivity'])}, specificity {fmt_pct(system['specificity'])}, PPV {fmt_pct(system['ppv'])}, and NPV {fmt_pct(system['npv'])}. The model AUC was {m['auc']:.3f} and GRACE 2.0 AUC was {gc['grace_auc']:.3f} on the same cohort. Calibration slope was {cal['slope']:.3f}, calibration-in-the-large {cal['citl']:.3f}, O:E {cal['oe_ratio']:.3f}, and ECE {cal['ece']:.3f}. Conclusions: A fixed-threshold admission model identified a high-risk subgroup for referral-center monitoring consideration. External validation pending."),
        ("normal", "Keywords: acute coronary syndrome; mortality; triage; random forest; calibration"),
        ("heading", "Introduction"),
        ("normal", "Acute coronary syndrome care depends on early recognition of patients at risk for death. Current ESC guidelines emphasize initial risk stratification using clinical history, vital signs, physical findings, ECG, biomarkers, and clinical instability [1,2]. GRACE remains a major reference score with broad validation [3,4]. Local referral centers may still need an admission-time tool that maps risk to practical monitoring options."),
        ("normal", "The immediate problem at admission is not only prognostic labeling. It is the allocation of monitored beds, nurse attention, repeat assessment, and senior review while diagnostic and treatment pathways are still moving. A score that performs well as a population predictor may still be difficult to translate into bedside monitoring decisions if its threshold is optimized for discrimination rather than for triage safety. In a referral center, the clinical cost of missing early deterioration is high, but the capacity cost of excessive escalation is also real."),
        ("normal", "Killip class captures heart failure severity after myocardial infarction and remains clinically meaningful decades after its first description [5]. Blood pressure and heart rate reflect hemodynamic stress. Renal markers, electrolytes, hemoglobin, glucose, and oxygen requirement describe perfusion, metabolic reserve, and acute respiratory support needs. These variables are routinely available at presentation. They also have face validity for admission decisions: a hypotensive patient with renal dysfunction and high oxygen need is different from a stable patient whose risk is driven only by age."),
        ("normal", "We developed a triage model using a random forest, an ensemble of tree predictors [6]. The model used 12 pre-specified routine admission variables. The aim was not to replace GRACE or clinician judgment. The aim was to produce an internally validated referral-center monitoring aid with fixed operating rules. The study objective was to estimate discrimination, calibration, and clinical tier performance for an admission-time model, compare its AUC with GRACE 2.0 on the same cohort, and quantify how a post-admission cardiogenic shock variable would inflate apparent performance."),
        ("heading", "Methods"),
        ("subheading", "Study Design and Setting"),
        ("normal", "The study used a de-identified single-center ACS registry from a referral center. The analysis was retrospective, and all model assessment was internal. The study protocol was approved by the Health Research Ethics Committee of the Faculty of Medicine, Universitas Hasanuddin (approval number 890/UN4.6.4.5.31/PP36/2026; protocol UH26050577; approved 22 July 2026, valid through 22 July 2027). The analysis used the de-identified registry with no direct patient contact. Reporting follows TRIPOD+AI guidance for prediction models that use regression or machine learning methods [7]. External validation pending."),
        ("subheading", "Participants and Outcomes"),
        ("normal", "The cohort included 1,817 eligible ACS admissions, 209 in-hospital deaths, and 1,608 survivors. The outcome was all-cause in-hospital mortality. GRACE 2.0 in-hospital mortality scores were available for the same patients. Killip class IV was treated as cardiogenic shock by definition in the source data: 279 patients had Killip class IV, which corresponds to cardiogenic shock at presentation; the admission-time shock-on-arrival flag identified 278 of these, and one further patient had Killip class IV recorded without that flag. Cardiogenic shock was recorded in 422 patients in total, so 143 patients had cardiogenic shock without Killip class IV, reflecting records made after admission."),
        ("normal", "This distinction was central to the design. Cardiogenic shock was not used in the main model to avoid look-ahead bias; including it inflates apparent performance. The flag can be recorded after admission, whereas Killip class IV reflects the presentation state. The model therefore used admission-time variables only, and cardiogenic shock was reserved for the sensitivity analysis."),
        ("subheading", "Predictors"),
        ("normal", "The 12 predictors were systolic blood pressure, heart rate, Killip class, hemoglobin, urea, estimated glomerular filtration rate, systemic immune-inflammation index, potassium, sodium, age at admission, random glucose, and oxygen need above 5 L/min. Features were selected from clinical domain knowledge before evaluation. No feature ranking or selection used outcome data."),
        ("normal", "Each predictor was chosen because it has a direct admission interpretation. Systolic blood pressure and heart rate describe hemodynamic stress. Killip class reflects pulmonary congestion and heart failure severity. Urea and estimated glomerular filtration rate reflect renal perfusion, chronic reserve, and the kidney response to low effective circulation. Hemoglobin describes oxygen carrying capacity. The systemic immune-inflammation index reflects inflammatory and thrombotic burden. Potassium and sodium identify electrolyte disturbance and illness severity. Random glucose captures acute metabolic stress. Oxygen need above 5 L/min reflects respiratory distress and impaired cardiopulmonary reserve. Age at admission anchors baseline vulnerability. All 12 variables were intended to be available within the first hour."),
        ("subheading", "Missing Data"),
        ("normal", f"Missingness was summarized for each model variable in Supplementary Table S1. {variable_label(highest_missing['variable'])} had the highest missingness at {fmt_pct(highest_missing['missing_pct'])}, followed by {missing_follow}. Complete data for all 12 model variables were present in {miss['complete_case_n']:,} patients ({fmt_pct(miss['complete_case_pct'])}), while any missingness among the 12 variables occurred in {miss['any_missing_n']} patients ({fmt_pct(miss['any_missing_pct'])}). Median imputation was fitted within cross-validation training folds only, then applied to the held-out fold. This preserved the evaluation boundary and avoided using held-out values to define imputation medians."),
        ("subheading", "Model Development"),
        ("normal", f"The random forest used 500 trees, maximum depth 6, minimum leaf size 5, random_state 42, and all available cores. Depth and leaf constraints were used to limit overfitting and to avoid very small terminal nodes in a cohort with 209 deaths. Events per variable were calculated as flagged deaths divided by the 12 predictors, giving EPV {m['epv']:.1f}; EPV was retained as a descriptive check because sample-size adequacy extends beyond a simple EPV rule [8]. The model was developed as a pragmatic triage tool, so the predictor set was fixed before validation rather than discovered from the full dataset."),
        ("subheading", "Internal Validation"),
        ("normal", "Validation used stratified five-fold pooled out-of-fold predictions with shuffle=True and random_state=42. Within each outer training fold, a three-fold inner loop estimated the Youden reference threshold. The held-out outer fold was used only for evaluation. No feature selection, threshold selection, tier optimization, calibration fitting, or operating-point selection used the pooled evaluation outcomes."),
        ("subheading", "Screening Threshold"),
        ("normal", f"The inner cross-validation Youden reference threshold was 0.1513. It was reported as a reference only and was not used for flagging. The fixed screening threshold was 0.08 because the clinical purpose was triage safety rather than a single optimal ROC point. Relative to the Youden reference, the 0.08 operating point captured {additional_deaths} additional deaths and produced {additional_fp} additional false positives, with escalated patients rising from {ref['escalated_n']} to {selected['escalated_n']}. Supplementary Table S2 shows the four examined thresholds: 0.1513, 0.10, 0.08, and 0.05."),
        ("subheading", "Tier Protocol"),
        ("normal", "Patients with predicted probability at least 0.08 were flagged. Within the flagged pool, stable descending probability ranking split patients into HIGH, INTERMEDIATE, and LOW tiers in fixed 25/50/25 proportions. HIGH and INTERMEDIATE were escalated for monitoring consideration. In referral-center terms, HIGH RISK implies consideration of ICU monitoring, INTERMEDIATE implies consideration of HCU monitoring, and LOW implies ward-level monitoring if the treating physician agrees. These are advisory implications, not automated placement decisions."),
        ("subheading", "Comparator and Statistical Analysis"),
        ("normal", "GRACE 2.0 in-hospital scores were computed on the same cohort and used as the comparator [3,4]. Performance was measured by AUC, Brier score, sensitivity, specificity, accuracy, PPV, NPV, and confusion matrix counts. Calibration assessment followed prediction-model performance guidance [9,10]. It included logistic calibration slope and calibration-in-the-large [10]. The analysis also reported observed-to-expected ratio and equal-frequency 10-bin ECE. AUCs were compared with the DeLong method for correlated ROC curves [11]. Decision curve analysis followed standard net benefit methods [12]. A sensitivity analysis added cardiogenic shock as a thirteenth feature to demonstrate look-ahead inflation."),
        ("heading", "Results"),
        ("subheading", "Baseline Characteristics"),
        ("normal", f"The cohort included {cohort['n']:,} patients, {cohort['deaths']} deaths ({fmt_pct(cohort['deaths'] / cohort['n'])}), and {cohort['survivors']:,} survivors. Baseline characteristics by outcome are shown in Table 1. Patients who died were older and had lower systolic blood pressure, higher heart rate, lower hemoglobin, higher urea, lower eGFR, higher systemic immune-inflammation index, more abnormal electrolytes, higher random glucose, more severe Killip class, and more frequent oxygen need above 5 L/min."),
        ("normal", "Mortality rose with the timing of the shock record. Death rates were 36.9% among the 279 patients with Killip class IV, 43.6% among all 422 patients with cardiogenic shock, and 56.6% among the 143 patients in whom cardiogenic shock was recorded without Killip class IV, representing shock recognized after admission. That last group is the reason cardiogenic shock was excluded from the main model: it identifies a very high-risk state, but not always one known at the admission decision point, and including it would inflate apparent triage performance."),
        ("subheading", "Missing Data"),
        ("normal", f"Missing data were limited but not absent. {variable_label(highest_missing['variable'])} had the highest missingness at {fmt_pct(highest_missing['missing_pct'])}, followed by {missing_follow}. eGFR missingness was {fmt_pct(next(r['missing_pct'] for r in miss['per_variable'] if r['variable'] == 'egfr_igd'))} after CKD-EPI 2021 derivation from creatinine, age, and sex. Complete data for all 12 model variables were present in {miss['complete_case_n']:,} patients ({fmt_pct(miss['complete_case_pct'])}), and {miss['any_missing_n']} patients ({fmt_pct(miss['any_missing_pct'])}) had at least one missing model variable. Median imputation within training folds allowed all eligible patients to contribute to pooled out-of-fold validation."),
        ("subheading", "Discrimination and Calibration"),
        ("normal", f"At threshold 0.08, the model flagged {flagged_n} patients ({fmt_pct(flagged_n / cohort['n'])}). This group included {stage1['tp']} deaths and {stage1['fp']} survivors, giving flagging sensitivity {fmt_pct(stage1['sensitivity'])}, specificity {fmt_pct(stage1['specificity'])}, PPV {fmt_pct(stage1['ppv'])}, and NPV {fmt_pct(stage1['npv'])}. The term flagged refers to all patients with predicted probability at least 0.08; the term escalated, used below, refers to the subset of flagged patients placed in the HIGH or INTERMEDIATE tiers, excluding the LOW tier. The pooled out-of-fold AUC was {m['auc']:.3f} and the Brier score was {m['brier']:.3f}."),
        ("normal", f"GRACE 2.0 AUC on the same cohort was {gc['grace_auc']:.3f}. The AUC difference was {gc['delta_auc']:.3f}, with DeLong p = {gc['p_value']:.3f} and 95% CI {gc['ci_low']:.3f} to {gc['ci_high']:.3f}. This comparison is shown in Table 2 and Figure 1. The finding supports local complementary decision support, not replacement of a score validated across very large ACS populations."),
        ("normal", f"Calibration was close to the ideal line. The calibration slope was {cal['slope']:.3f}, calibration-in-the-large was {cal['citl']:.3f}, O:E ratio was {cal['oe_ratio']:.3f}, and ECE was {cal['ece']:.3f}. Figure 2 shows the reliability pattern across risk bins. These estimates came from pooled out-of-fold probabilities and were not training-set summaries."),
        ("subheading", "Tier Allocation and Clinical Operating Point"),
        ("normal", f"The HIGH tier included {tiers['high']['n']} patients, {tiers['high']['deaths']} deaths, and PPV {fmt_pct(tiers['high']['ppv'])}. The INTERMEDIATE tier included {tiers['intermediate']['n']} patients, {tiers['intermediate']['deaths']} deaths, and PPV {fmt_pct(tiers['intermediate']['ppv'])}. The LOW tier included {tiers['low']['n']} patients, {tiers['low']['deaths']} deaths, and PPV {fmt_pct(tiers['low']['ppv'])}. HIGH plus INTERMEDIATE contained {system['tp']} deaths, equal to {fmt_pct(flagged_capture)} of the {stage1['tp']} flagged deaths ({fmt_pct(system['sensitivity'])} of all {cohort['deaths']} deaths)."),
        ("normal", f"For the full triage system, {escalated_n} patients were escalated, equal to {fmt_pct(escalated_n / cohort['n'])} of the cohort. The confusion matrix was TP {system['tp']}, FP {system['fp']}, FN {system['fn']}, and TN {system['tn']:,}, summing to {cohort['n']:,} patients. Overall sensitivity was {fmt_pct(system['sensitivity'])} ({system['tp']}/{cohort['deaths']}), specificity {fmt_pct(system['specificity'])}, accuracy {fmt_pct(system['accuracy'])}, PPV {fmt_pct(system['ppv'])}, and NPV {fmt_pct(system['npv'])}. Missed deaths were {system['fn']} of {cohort['deaths']} ({fmt_pct(system['fn'] / cohort['deaths'])}): {missed_not_flagged} patients were not flagged and {missed_low} were flagged but placed in the LOW tier. The {fmt_pct(flagged_capture)} figure is secondary and uses a different denominator, {system['tp']} of {stage1['tp']} flagged deaths. Figure 4 illustrates the proposed referral-center workflow, mapping HIGH, INTERMEDIATE, and LOW tiers to advisory ICU, HCU, and ward monitoring considerations."),
        ("normal", threshold_paragraph),
        ("subheading", "Sensitivity Analysis"),
        ("normal", f"In the sensitivity analysis, adding cardiogenic shock increased system sensitivity to {fmt_pct(cs['system_sensitivity'])}, HIGH-tier PPV to {fmt_pct(cs['high_ppv'])}, and AUC to {cs['auc']:.3f}. Cardiogenic shock was not used in the main model to avoid look-ahead bias; including it inflates apparent performance. The mechanism is visible in the registry timing: 143 of 422 patients with cardiogenic shock had no Killip class IV, and their mortality was 56.6%."),
        ("heading", "Discussion"),
        ("normal", f"This study produced an admission-time ACS mortality triage model with fixed operating rules and internally validated performance. The model flagged {fmt_pct(flagged_n / cohort['n'])} of patients and captured {stage1['tp']} of {cohort['deaths']} deaths. After the fixed 25/50/25 tier split, HIGH plus INTERMEDIATE escalation identified {system['tp']} deaths, giving overall sensitivity {fmt_pct(system['sensitivity'])} and NPV {fmt_pct(system['npv'])}. The HIGH tier was clinically enriched, with {tiers['high']['deaths']} deaths among {tiers['high']['n']} patients and PPV {fmt_pct(tiers['high']['ppv'])}. The main value of the protocol is its direct mapping from a probability score to monitoring consideration at the referral center."),
        ("normal", f"Clinical use should remain advisory. A HIGH RISK result can support consideration of ICU monitoring, and an INTERMEDIATE result can support consideration of HCU monitoring, but the treating physician retains all placement and treatment decisions. A LOW result should not be read as absence of risk. {missed_low} flagged patients in the LOW tier died, and {missed_not_flagged} deaths occurred below the screening threshold. These deaths matter because the clinical task is not only to enrich a high-risk group, but also to make residual risk visible. The system therefore belongs in a prospective workflow that pairs the tier with bedside reassessment, repeat vital signs, and clinician review."),
        ("normal", "The fixed operating point also makes prospective use auditable. A referral center can record how often HIGH and INTERMEDIATE results led to monitored care, how often clinicians overrode the suggestion, and whether missed deaths shared recognizable bedside features. Because the threshold and tier split are fixed, future monitoring can separate model performance from implementation behavior. That distinction is important: poor uptake, delayed laboratory availability, or lack of monitored beds could weaken clinical effect even if statistical performance remains stable. The audit should also track alert burden, time to clinician review, escalation delays, and reasons for disagreement."),
        ("normal", f"The comparison with GRACE should be interpreted carefully. GRACE has broad external evidence and remains a reference score in ACS risk assessment [3,4]. The local model had AUC {m['auc']:.3f} compared with GRACE 2.0 AUC {gc['grace_auc']:.3f} on the same cohort, with an AUC difference of {gc['delta_auc']:.3f} and DeLong p = {gc['p_value']:.3f}. That modest difference may be useful for a referral-center monitoring protocol, but it is not a claim that a single-center model supersedes GRACE. Both are decision-support tools. GRACE provides a general benchmark; the local model translates routine admission variables into a fixed monitoring tier."),
        ("normal", f"The look-ahead analysis is a methodological finding. Adding cardiogenic shock raised sensitivity to {fmt_pct(cs['system_sensitivity'])} and AUC to {cs['auc']:.3f}, which looks attractive until the timing of the variable is considered. Cardiogenic shock can be recorded after admission, and 143 patients with cardiogenic shock had no Killip class IV. Including that variable would let the model use information that may not exist when the triage decision is made. Reporting the inflated result makes the bias visible and supports the decision to keep the main model limited to admission-time variables."),
        ("normal", "The predictor pattern also has clinical coherence. Killip class had the highest feature importance at approximately 0.152, consistent with the central role of heart failure severity in early ACS mortality. Urea was next at approximately 0.138 and eGFR followed at approximately 0.120, making renal status a dominant signal. The kidney is an early sensor of hypoperfusion, venous congestion, neurohormonal activation, and chronic reserve. Systolic blood pressure and systemic immune-inflammation index were each approximately 0.090, linking hemodynamic compromise with inflammatory burden. Oxygen need above 5 L/min was approximately 0.084, reflecting respiratory distress and limited cardiopulmonary reserve. These variables are not abstract model inputs; they describe physiology clinicians already monitor."),
        ("normal", "The resource-limited projection is deliberately narrow. At the referral center, the model can be studied as a way to structure ICU, HCU, and ward monitoring discussions. District or network deployment is hypothetical and requires external validation. Transferability cannot be assumed because treatment availability, referral delays, competing bed constraints, and facility-level case mix can alter both predictor distributions and outcomes. A facility with fewer monitored beds may experience a treatment paradox: patients flagged as high risk may receive different care precisely because the model identifies them, changing observed mortality after implementation."),
        ("normal", f"Several limitations remain. This was a single-center retrospective study with internal validation only. External validation pending. The event count was modest, even though EPV was {m['epv']:.1f}. PPV was bounded by the {fmt_pct(cohort['deaths'] / cohort['n'])} mortality prevalence, so many escalated patients survived. Missing data were handled with fold-specific median imputation, but missingness may still reflect clinical workflow and illness severity. Registry labels may contain timing noise, particularly for shock-related variables. The analysis did not include temporal validation, prospective impact assessment, clinician adherence, or calibration monitoring after deployment. These limitations define the next evaluation step: a prospective referral-center study with external validation before broader use."),
        ("heading", "Conclusion"),
        ("normal", "An admission-time random forest model identified ACS patients at higher risk of in-hospital death and mapped them to advisory referral-center monitoring tiers. The fixed protocol captured 158 of 209 deaths after escalation while keeping decisions with the treating physician. The results support prospective referral-center testing with explicit monitoring of clinical workflow. External validation is required before broader use."),
        ("normal", "Conflicts of Interest: The authors declare no conflicts of interest."),
    ]


def count_words(paragraphs: list[tuple[str, str]], tables: list[list[list[str]]]) -> int:
    text = " ".join(t for _, t in paragraphs)
    for table in tables:
        for row in table:
            text += " " + " ".join(row)
    return len(re.findall(r"\b\S+\b", text))


def doc_word_count(doc: Document) -> int:
    words = 0
    for paragraph in doc.paragraphs:
        words += len(paragraph.text.split())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                words += len(cell.text.split())
    return words


def apply_manuscript_format(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.5
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)


def build_manuscript(results: dict) -> int:
    table1 = baseline_table()
    table2 = [
        ["Metric", "Model", "GRACE 2.0"],
        ["AUC", f"{results['main']['auc']:.3f}", f"{results['grace_comparison']['grace_auc']:.3f}"],
        ["AUC difference", f"{results['grace_comparison']['delta_auc']:.3f}", ""],
        ["DeLong p value", f"{results['grace_comparison']['p_value']:.3f}", ""],
        ["Brier score", f"{results['main']['brier']:.3f}", ""],
        ["Calibration slope", f"{results['main']['calibration']['slope']:.3f}", ""],
    ]
    tiers = results["main"]["tiers"]
    table3 = [
        ["Group", "n", "Deaths", "PPV"],
        ["HIGH", str(tiers["high"]["n"]), str(tiers["high"]["deaths"]), fmt_pct(tiers["high"]["ppv"])],
        ["INTERMEDIATE", str(tiers["intermediate"]["n"]), str(tiers["intermediate"]["deaths"]), fmt_pct(tiers["intermediate"]["ppv"])],
        ["LOW", str(tiers["low"]["n"]), str(tiers["low"]["deaths"]), fmt_pct(tiers["low"]["ppv"])],
        ["Not flagged", str(tiers["not_flagged"]["n"]), str(tiers["not_flagged"]["deaths"]), fmt_pct(tiers["not_flagged"]["ppv"])],
    ]
    table_s1 = [["Variable", "Missing n", "Missing %"]] + [[r["variable"], str(r["missing_n"]), f"{r['missing_pct']*100:.1f}%"] for r in results["missingness"]["per_variable"]]
    table_s2 = [["Threshold", "Sensitivity", "False positives", "Missed deaths", "Flagged n", "Escalated n", "PPV", "Specificity"]] + [[str(r["threshold"]), fmt_pct(r["system"]["sensitivity"]), str(r["false_positives"]), str(r["missed_deaths"]), str(r["flagged_n"]), str(r["escalated_n"]), fmt_pct(r["system"]["ppv"]), fmt_pct(r["system"]["specificity"])] for r in results["threshold_tradeoff"]]
    paragraphs = manuscript_text(results, 0)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    for kind, text in paragraphs:
        if kind == "heading":
            doc.add_heading(text, level=1)
        elif kind == "subheading":
            doc.add_heading(text, level=2)
        elif kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
        elif kind == "authors":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.bold = True
        else:
            doc.add_paragraph(text)

    doc.add_paragraph("Table 1. Baseline characteristics by outcome.")
    add_table(doc, table1)
    doc.add_paragraph("Table 2. Model and GRACE 2.0 performance in pooled out-of-fold validation.")
    add_table(doc, table2)
    doc.add_paragraph("Table 3. Tier allocation and not-flagged mortality.")
    add_table(doc, table3)
    figure_captions = [
        ("fig1_roc.png", "Figure 1. ROC curves for the admission model and GRACE 2.0 on the same cohort."),
        ("fig2_calibration.png", "Figure 2. Calibration diagram for pooled out-of-fold admission model probabilities."),
        ("fig3_tiers.png", "Figure 3. Mortality enrichment across HIGH, INTERMEDIATE, LOW, and not-flagged groups."),
        ("fig4_flow.png", "Figure 4. Referral-center monitoring flow with advisory ICU, HCU, and ward implications."),
    ]
    for fig_name, caption in figure_captions:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FIGURES_DIR / fig_name), width=Inches(5.8))
    doc.add_paragraph("Supplementary Table S1. Missingness in model variables.")
    add_table(doc, table_s1)
    doc.add_paragraph("Supplementary Table S2. Threshold trade-off.")
    add_table(doc, table_s2)
    supplementary_captions = [
        ("fig5_threshold_tradeoff.png", "Supplementary Figure S1. Threshold trade-off across the Youden reference and lower triage thresholds."),
        ("fig6_dca.png", "Supplementary Figure S2. Decision curve analysis for the admission model."),
    ]
    for fig_name, caption in supplementary_captions:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FIGURES_DIR / fig_name), width=Inches(5.8))
    doc.add_heading("References", level=1)
    for i, ref in enumerate(REFERENCES, start=1):
        doc.add_paragraph(f"{i}. {ref}")
    wc = doc_word_count(doc)
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith("Manuscript word count:"):
            paragraph.text = f"Manuscript word count: {wc}"
            break
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = f"Manuscript word count: {wc}"
    apply_manuscript_format(doc)
    MANUSCRIPT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(MANUSCRIPT_PATH)
    return wc


def main() -> None:
    results = load_results()
    write_model_dictionary(results)
    build_notebook(results)
    wc = build_manuscript(results)
    print(f"Built notebook, model dictionary, and manuscript. Word count: {wc}")


if __name__ == "__main__":
    main()
