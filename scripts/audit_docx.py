#!/usr/bin/env python3
"""Independent audit of the generated manuscript docx."""
import re
import zipfile
import sys
from pathlib import Path

import docx

path = Path("manuscript/acs_mortality_triage_manuscript.docx")
doc = docx.Document(path)

# 1. Word count: paragraphs + table cells (split on whitespace)
words = 0
for p in doc.paragraphs:
    words += len(p.text.split())
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            words += len(cell.text.split())
print(f"WORD COUNT (paragraphs + table cells): {words}")

# 2. Full text including tables
texts = [p.text for p in doc.paragraphs]
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            texts.append(cell.text)
full = "\n".join(texts)

# 3. Em dashes / en dashes
print(f"EM DASHES (—): {full.count(chr(8212))}")
print(f"EN DASHES (–): {full.count(chr(8211))}")

# 4. Forbidden terms
for term in ["echo", "lvef", "tapse", "lvot", "non-echocardiographic", "echocardiograph"]:
    hits = len(re.findall(term, full, re.I))
    if hits:
        print(f"FORBIDDEN '{term}': {hits} hits")

# 5. Delta/iteration language
for phrase in ["earlier analysis", "previous version", "in contrast to previous", "we improved", "unlike"]:
    hits = len(re.findall(phrase, full, re.I))
    if hits:
        print(f"DELTA-LANG '{phrase}': {hits}")

# 6. Key numbers present
for num in ["1,817", "209", "691", "158", "359", "51", "1,249", "75.6", "50.6", "0.08", "cardiogenic shock"]:
    print(f"number '{num}': {'OK' if num in full else 'MISSING'}")

# 7. References count (Vancouver list at end)
refs = re.findall(r"^\s*\d+\.\s", full, re.M)
print(f"REFERENCE-LIST-LIKE LINES: {len(refs)}")

# 8. Figures embedded?
with zipfile.ZipFile(path) as z:
    media = [n for n in z.namelist() if n.startswith("word/media/")]
    print(f"EMBEDDED MEDIA: {media}")

# 9. Images referenced in body
import docx.oxml.ns as ns
body = doc.element.body
drawings = body.findall(".//" + ns.qn("w:drawing"))
print(f"DRAWING ELEMENTS: {len(drawings)}")

# 10. Headings present
heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
print("HEADINGS:", heads[:25])
