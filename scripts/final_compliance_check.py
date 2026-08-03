#!/usr/bin/env python3
"""Final compliance check: manuscript numbers vs latest run JSON,
Vancouver citation order, reference list vs verified-references.md."""
import json
import re
import sys

import docx

d = json.load(open("results/analysis_results.json"))
doc = docx.Document("manuscript/acs_mortality_triage_manuscript.docx")
paras = [p.text for p in doc.paragraphs]
full = "\n".join(paras)
for t in doc.tables:
    for r in t.rows:
        for c in r.cells:
            full += "\n" + c.text

m = d["main"]
tiers = m["tiers"]
sys_ = m["system"]
gc = d["grace_comparison"]
cal = m["calibration"]
miss = d["missingness"]
pv = {x["variable"]: x for x in miss["per_variable"]}
soa = d["shock_crosstabs"]
trade = {r["threshold"]: r for r in d["threshold_tradeoff"]}
withcs = d["sensitivity_analysis"]["with_cardiogenic_shock_13_features"]
n = d["cohort"]["n"]
deaths = d["cohort"]["deaths"]

exp = {
    # (label, must_appear)
    "flagged 682": f"{m['stage1']['tp'] + m['stage1']['fp']:,}",
    "flag deaths 172": f"{m['stage1']['tp']}",
    "HIGH 170/88": f"{tiers['high']['n']}",
    "INT 341": f"{tiers['intermediate']['n']}",
    "LOW 171": f"{tiers['low']['n']}",
    "escalated 511": f"{tiers['high']['n'] + tiers['intermediate']['n']}",
    "CM tp 158": f"{sys_['tp']}",
    "CM fp 353": f"{sys_['fp']}",
    "CM tn 1255": f"{sys_['tn']:,}",
    "sens 75.6": f"{sys_['sensitivity'] * 100:.1f}%",
    "spec 78.0": f"{sys_['specificity'] * 100:.1f}%",
    "acc 77.8": f"{sys_['accuracy'] * 100:.1f}%",
    "PPV 30.9": f"{sys_['ppv'] * 100:.1f}%",
    "NPV 96.1": f"{sys_['npv'] * 100:.1f}%",
    "AUC 0.843": f"{m['auc']:.3f}",
    "GRACE 0.816": f"{gc['grace_auc']:.3f}",
    "delta 0.027": f"{gc['delta_auc']:.3f}",
    "cal slope 1.076": f"{cal['slope']:.3f}",
    "cal citl 0.015": f"{cal['citl']:.3f}",
    "cal oe 1.014": f"{cal['oe_ratio']:.3f}",
    "cal ece 0.017": f"{cal['ece']:.3f}",
    "withCS sens 90.4": f"{withcs['system_sensitivity'] * 100:.1f}%",
    "withCS high ppv 63.2": f"{withcs['high_ppv'] * 100:.1f}%",
    "withCS auc 0.938": f"{withcs['auc']:.3f}",
    "Killip IV 279": f"{soa['killip_iv_n']}",
    "CS 422": f"{soa['cardiogenic_shock_n']}",
    "CS-no-KIV 143": f"{soa['cs_without_killip_iv_n']}",
    "urea missing 1.4": f"{pv['ureum_igd']['missing_pct'] * 100:.1f}%",
    "eGFR missing 0.6": f"{pv['egfr_igd']['missing_pct'] * 100:.1f}%",
    "complete 1,745": f"{miss['complete_case_n']:,}",
    "any 72": f"{miss['any_missing_n']}",
    "trade 0.1513 escal 327": f"{trade[0.1513]['escalated_n']}",
    "trade 0.10 escal 439": f"{trade[0.10]['escalated_n']}",
    "trade 0.05 escal 670": f"{trade[0.05]['escalated_n']}",
}
print("=== ANGKA vs JSON (harus SEMUA OK) ===")
bad = 0
for label, val in exp.items():
    # flexible match: value may appear with different punctuation
    pat = re.escape(val)
    ok = re.search(pat, full) is not None
    if not ok:
        bad += 1
        print(f"  MISSING [{label}] = '{val}' TIDAK ADA di docx")
print(f"  {len(exp) - bad}/{len(exp)} cocok")

stale = ["691", "517", "1249", "6.7%", "50.6%", "77.7%", "0.842", "1.088", "56.2%", "37.1%", "p = 0.026", "0.025"]
print("\n=== ANGKA STALE (harus NOL) ===")
found_stale = 0
for s in stale:
    hits = [l.strip()[:100] for l in full.split("\n") if s in l]
    if hits:
        found_stale += 1
        print(f"  STALE '{s}': {hits[0]}")
print(f"  {found_stale} nilai stale")

print("\n=== EPV (harus 14.3 = 172/12) ===")
print(f"  EPV JSON: {m['epv']:.1f}; docx menyebut: {[l.strip()[:80] for l in full.split(chr(10)) if 'EPV' in l][:1]}")

print("\n=== TAGS IN-TEXT vs URUTAN VANCOUVER ===")
tags = re.findall(r"\[(\d+(?:,\s*\d+)*)\]", full)
seen, first_order = set(), []
for t in tags:
    for num in re.split(r",\s*", t):
        n = int(num)
        if n not in seen:
            seen.add(n)
            first_order.append(n)
print(f"  urutan kemunculan pertama: {first_order}")
ok_order = first_order == list(range(1, 13))
print(f"  urutan 1..12 sempurna: {ok_order}")
max_tag = max(seen)
print(f"  max tag: {max_tag} (harus 12)")

print("\n=== DAFTAR REFERENSI DOCX vs verified-references.md ===")
ref_start = next(i for i, t in enumerate(paras) if t.strip() == "References")
docx_refs = [t.strip() for t in paras[ref_start + 1:] if t.strip()]
verified = [l.strip() for l in open("references/verified-references.md") if re.match(r"^\d+\.\s", l.strip())]
print(f"  docx refs: {len(docx_refs)} | verified: {len(verified)}")
for i, (dr, vr) in enumerate(zip(docx_refs, verified), 1):
    vnum = re.sub(r"^\d+\.\s*", "", vr)
    match = vnum.split(".")[0][:40].lower() in dr.lower() or dr.split(".")[0][:40].lower() in vnum.lower()
    print(f"  [{i}] {'OK ' if match else 'CEK'} docx: {dr[:70]}")
    if not match:
        print(f"        verified: {vnum[:90]}")
