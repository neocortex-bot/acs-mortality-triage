"""Build a journal submission package (single zip) from analysis results.

OJS-style components:
  01_Article_Text.docx   - anonymous body, single-spaced, 12pt, NO embedded
                           tables/figures (separated per journal requirements)
  02_Title_Page.docx     - title, authors, affiliations, ORCID, corresponding,
                           running title, word count, declarations
  03_Cover_Letter.docx   - cover letter
  04_Tables/Table_{1,2,3,S1,S2}.docx  - each with caption + table
  05_Figures/Figure_{1..4}.tiff       - 300 DPI TIFF
  05_Figures/Supplementary_Figure_{S1,S2}.tiff
  05_Figures/Figure_Legends.docx
  UPLOAD_GUIDE.txt       - mapping to OJS components

Reuses builders from scripts/build_artifacts.py (data-driven, no hardcoded
metrics).
"""

from __future__ import annotations

import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from scripts.build_artifacts import (
    REFERENCES,
    add_table,
    baseline_table,
    fmt_pct,
    load_results,
    manuscript_text,
)

SUBMISSION_DIR = ROOT / "submission"
ZIP_PATH = ROOT / "submission" / "acs_mortality_triage_submission_package.zip"

TITLE = "A Pragmatically Calibrated Machine Learning Triage System for In-Hospital Mortality in Acute Coronary Syndromes"

FIGURE_LEGENDS = {
    "fig1_roc.png": "Figure 1. ROC curves for the admission model and GRACE 2.0 on the same cohort.",
    "fig2_calibration.png": "Figure 2. Calibration diagram for cross-validated admission model probabilities.",
    "fig3_tiers.png": "Figure 3. Mortality enrichment across HIGH, INTERMEDIATE, LOW, and not-flagged groups.",
    "fig4_flow.png": "Figure 4. Referral-center monitoring flow with advisory ICU, HCU, and ward implications.",
    "fig5_threshold_tradeoff.png": "Supplementary Figure S1. Threshold trade-off across the Youden reference and lower triage thresholds.",
    "fig6_dca.png": "Supplementary Figure S2. Decision curve analysis for the admission model.",
}

TABLE_CAPTIONS = {
    "Table_1": "Table 1. Baseline characteristics by outcome.",
    "Table_2": "Table 2. Model and GRACE 2.0 performance in cross-validated analysis.",
    "Table_3": "Table 3. Tier allocation and not-flagged mortality.",
    "Table_S1": "Supplementary Table S1. Missingness in model variables.",
    "Table_S2": "Supplementary Table S2. Threshold trade-off.",
}


def _style_doc(doc: Document, single_spaced: bool = True) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.0 if single_spaced else 1.5


def _add_heading(doc: Document, text: str, size: int = 12) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def _add_centered(doc: Document, text: str, bold: bool = False, size: int = 12) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def build_article_text() -> Path:
    """Anonymous body: no authors, no embedded tables/figures, single-spaced."""
    results = load_results()
    paragraphs = manuscript_text(results)
    doc = Document()
    _style_doc(doc, single_spaced=True)

    skip_prefixes = (
        "Manuscript word count:",
        "Running title:",
        "Corresponding author:",
        "ORCID:",
    )
    for kind, text in paragraphs:
        if kind == "authors":
            continue
        if text.startswith("\u00b9 ") or text.startswith("\u00b2 "):
            continue  # affiliation lines
        if text.startswith(skip_prefixes):
            continue
        if kind == "title":
            _add_centered(doc, text, bold=True, size=14)
        elif kind == "heading":
            _add_heading(doc, text, size=13)
        elif kind == "subheading":
            _add_heading(doc, text, size=12)
        else:
            doc.add_paragraph(text)

    path = SUBMISSION_DIR / "01_Article_Text.docx"
    doc.save(path)
    return path


def build_title_page() -> Path:
    doc = Document()
    _style_doc(doc, single_spaced=True)
    _add_centered(doc, TITLE, bold=True, size=14)
    doc.add_paragraph()
    _add_centered(doc, "Izzan Rijal Muslim, MD\u00b9; Andriany Qanitha, MD, PhD, FESC\u00b2", bold=True)
    doc.add_paragraph("\u00b9 Department of Cardiology and Vascular Medicine, Faculty of Medicine, Hasanuddin University, Makassar, Indonesia")
    doc.add_paragraph("\u00b2 Department of Physiology, Clinical Epidemiology, Research, and Publication Unit, Faculty of Medicine, Hasanuddin University, Makassar, Indonesia")
    doc.add_paragraph("ORCID: Izzan Rijal Muslim 0009-0005-0173-3606; Andriany Qanitha 0000-0003-2420-0560")
    doc.add_paragraph("Corresponding author: Izzan Rijal Muslim, MD. Email: izzan.rijal@gmail.com")
    doc.add_paragraph("Running title: ACS triage")
    doc.add_paragraph("Keywords: acute coronary syndrome; mortality; triage; random forest; calibration")
    doc.add_paragraph()
    _add_heading(doc, "Declarations")
    doc.add_paragraph("Ethics approval: The study protocol was approved by the Health Research Ethics Committee of the Faculty of Medicine, Universitas Hasanuddin (approval number 890/UN4.6.4.5.31/PP36/2026). The analysis used the de-identified registry with no direct patient contact.")
    doc.add_paragraph("Conflicts of interest: The authors declare no conflicts of interest.")
    doc.add_paragraph("Funding: None.")
    path = SUBMISSION_DIR / "02_Title_Page.docx"
    doc.save(path)
    return path


def build_cover_letter() -> Path:
    doc = Document()
    _style_doc(doc, single_spaced=True)
    doc.add_paragraph("August 3, 2026")
    doc.add_paragraph()
    doc.add_paragraph("The Editor-in-Chief")
    doc.add_paragraph("[Journal name and address]")
    doc.add_paragraph()
    doc.add_paragraph('Re: Submission of original article \u2014 "A Pragmatically Calibrated Machine Learning Triage System for In-Hospital Mortality in Acute Coronary Syndromes"')
    doc.add_paragraph()
    doc.add_paragraph("Dear Editor,")
    doc.add_paragraph()
    doc.add_paragraph('We wish to submit our original research article, entitled "A Pragmatically Calibrated Machine Learning Triage System for In-Hospital Mortality in Acute Coronary Syndromes," for consideration for publication in [Journal Name].')
    doc.add_paragraph()
    doc.add_paragraph("The study develops and internally validates an admission-time triage model for patients with acute coronary syndrome, using twelve routine clinical variables that are available within the first hour of presentation. In a registry of 1,817 patients, the model identified a high-risk subgroup for referral-center monitoring consideration, with a prespecified threshold and tier protocol designed for practical use in settings with limited monitored-bed capacity. Its discrimination was comparable to GRACE 2.0 on the same cohort (AUC 0.843 versus 0.816), and we report calibration, decision-curve, and threshold trade-off analyses. The study is reported in accordance with the TRIPOD+AI statement.")
    doc.add_paragraph()
    doc.add_paragraph("We believe the manuscript fits the journal's scope because it addresses a practical problem in resource-limited cardiovascular care: translating a probability score into monitoring decisions (ICU, high-care, or ward) rather than only labeling risk. All analyses used de-identified registry data, and the study was approved by the Health Research Ethics Committee of the Faculty of Medicine, Universitas Hasanuddin (approval number 890/UN4.6.4.5.31/PP36/2026).")
    doc.add_paragraph()
    doc.add_paragraph("This manuscript is original, has not been published previously, and is not under consideration by any other journal. All authors have read and approved the final version and agree to its submission. The authors declare no conflicts of interest.")
    doc.add_paragraph()
    doc.add_paragraph("We thank you for considering our work.")
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph("Izzan Rijal Muslim, MD")
    doc.add_paragraph("Department of Cardiology and Vascular Medicine, Faculty of Medicine, Hasanuddin University, Makassar, Indonesia")
    doc.add_paragraph("Email: izzan.rijal@gmail.com | ORCID: 0009-0005-0173-3606")
    doc.add_paragraph("On behalf of all authors")
    path = SUBMISSION_DIR / "03_Cover_Letter.docx"
    doc.save(path)
    return path


def build_tables(results: dict) -> list[Path]:
    table1 = baseline_table()
    table2 = [
        ["Metric", "Model", "GRACE 2.0"],
        ["AUC", f"{results['main']['auc']:.3f}", f"{results['grace_comparison']['grace_auc']:.3f}"],
        ["AUC difference", f"{results['grace_comparison']['delta_auc']:.3f}", ""],
        ["DeLong p value", f"{results['grace_comparison']['p_value']:.3f}", ""],
        ["Brier score", f"{results['main']['brier']:.3f}", ""],
        ["Calibration slope", f"{results['main']['calibration']['slope']:.3f}", ""],
    ]
    tiers = results["main"]["tiers"]
    table3 = [
        ["Group", "n", "Deaths", "PPV"],
        ["HIGH", str(tiers["high"]["n"]), str(tiers["high"]["deaths"]), fmt_pct(tiers["high"]["ppv"])],
        ["INTERMEDIATE", str(tiers["intermediate"]["n"]), str(tiers["intermediate"]["deaths"]), fmt_pct(tiers["intermediate"]["ppv"])],
        ["LOW", str(tiers["low"]["n"]), str(tiers["low"]["deaths"]), fmt_pct(tiers["low"]["ppv"])],
        ["Not flagged", str(tiers["not_flagged"]["n"]), str(tiers["not_flagged"]["deaths"]), fmt_pct(tiers["not_flagged"]["ppv"])],
    ]
    table_s1 = [["Variable", "Missing n", "Missing %"]] + [
        [r["variable"], str(r["missing_n"]), f"{r['missing_pct']*100:.1f}%"]
        for r in results["missingness"]["per_variable"]
    ]
    table_s2 = [["Threshold", "Sensitivity", "False positives", "Missed deaths", "Flagged n", "Escalated n", "PPV", "Specificity"]] + [
        [str(r["threshold"]), fmt_pct(r["system"]["sensitivity"]), str(r["false_positives"]),
         str(r["missed_deaths"]), str(r["flagged_n"]), str(r["escalated_n"]),
         fmt_pct(r["system"]["ppv"]), fmt_pct(r["system"]["specificity"])]
        for r in results["threshold_tradeoff"]
    ]
    tables = {"Table_1": table1, "Table_2": table2, "Table_3": table3, "Table_S1": table_s1, "Table_S2": table_s2}
    paths = []
    for name, rows in tables.items():
        doc = Document()
        _style_doc(doc, single_spaced=True)
        doc.add_paragraph(TABLE_CAPTIONS[name])
        add_table(doc, rows)
        path = SUBMISSION_DIR / "04_Tables" / f"{name}.docx"
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(path)
        paths.append(path)
    return paths


def build_figures() -> list[Path]:
    from src.config import FIGURES_DIR
    # Map source PNG -> final OJS-style file name
    name_map = {
        "fig1_roc.png": "Figure_1.tiff",
        "fig2_calibration.png": "Figure_2.tiff",
        "fig3_tiers.png": "Figure_3.tiff",
        "fig4_flow.png": "Figure_4.tiff",
        "fig5_threshold_tradeoff.png": "Supplementary_Figure_S1.tiff",
        "fig6_dca.png": "Supplementary_Figure_S2.tiff",
    }
    paths = []
    for png_name, caption in FIGURE_LEGENDS.items():
        src = FIGURES_DIR / png_name
        if not src.exists():
            continue
        img = Image.open(src).convert("RGB")
        path = SUBMISSION_DIR / "05_Figures" / name_map[png_name]
        path.parent.mkdir(parents=True, exist_ok=True)
        img.save(path, format="TIFF", compression="tiff_lzw", dpi=(300, 300))
        paths.append(path)
    return paths


def build_figure_legends() -> Path:
    doc = Document()
    _style_doc(doc, single_spaced=True)
    _add_heading(doc, "Figure Legends")
    for caption in FIGURE_LEGENDS.values():
        doc.add_paragraph(caption)
    path = SUBMISSION_DIR / "05_Figures" / "Figure_Legends.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def build_upload_guide() -> Path:
    guide = """UPLOAD GUIDE - OJS SUBMISSION
============================

Upload each file under the matching OJS "Article Component":

1. Upload File step, component "Article Text"
   -> 01_Article_Text.docx
      (anonymous body; single-spaced, 12 pt Times New Roman; tables and
       figures are NOT embedded - they are separate files per journal
       requirements. In-text callouts remain, e.g. "Table 1", "Figure 1".)

2. Component "Title page"
   -> 02_Title_Page.docx
      (title, authors, affiliations, ORCID, corresponding author, running
       title, keywords, declarations)

3. Component "Cover Letter"
   -> 03_Cover_Letter.docx
      (replace [Journal name and address] and [Journal Name] with the
       journal details before uploading)

4. Component "Tables" (one file per table, .docx, each with its caption)
   -> 04_Tables/Table_1.docx
   -> 04_Tables/Table_2.docx
   -> 04_Tables/Table_3.docx
   -> 04_Tables/Table_S1.docx
   -> 04_Tables/Table_S2.docx

5. Component "Figures" (TIFF, 300 DPI, with captions in Figure_Legends.docx)
   -> 05_Figures/Figure_1.tiff   (ROC curves)
   -> 05_Figures/Figure_2.tiff   (calibration diagram)
   -> 05_Figures/Figure_3.tiff   (mortality enrichment)
   -> 05_Figures/Figure_4.tiff   (monitoring flow)
   -> 05_Figures/Figure_Legends.docx

6. Component "Supplementary File"
   -> 05_Figures/Supplementary_Figure_S1.tiff  (threshold trade-off)
   -> 05_Figures/Supplementary_Figure_S2.tiff  (decision curve analysis)
   -> 04_Tables/Table_S1.docx and Table_S2.docx may also be listed here
      if the journal prefers supplementary items grouped together.

Notes
-----
* Word count (text only, including references): computed from
  01_Article_Text.docx at build time; the canonical full count including
  table text is 3,620.
* References are Vancouver style with DOIs (URLs available via doi.org).
* Ethics approval number is stated on the title page and in the body.
* If the journal is NOT double-blind, author details are already present in
  02_Title_Page.docx; you may also merge them into 01_Article_Text.docx.
"""
    path = SUBMISSION_DIR / "UPLOAD_GUIDE.txt"
    path.write_text(guide, encoding="utf-8")
    return path


def main() -> None:
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)
    results = load_results()
    files = [
        build_article_text(),
        build_title_page(),
        build_cover_letter(),
        *build_tables(results),
        *build_figures(),
        build_figure_legends(),
        build_upload_guide(),
    ]
    # Word count of article text (text only) for the guide
    from docx import Document as D
    art = D(SUBMISSION_DIR / "01_Article_Text.docx")
    wc = sum(len(p.text.split()) for p in art.paragraphs)
    guide = SUBMISSION_DIR / "UPLOAD_GUIDE.txt"
    guide_text = guide.read_text(encoding="utf-8")
    old_line = "* Word count (text only, including references): computed from\n  01_Article_Text.docx at build time;"
    new_line = f"* Word count (text only, including references): {wc};"
    assert old_line in guide_text, "guide placeholder not found"
    guide.write_text(guide_text.replace(old_line, new_line), encoding="utf-8")

    if ZIP_PATH.exists():
        ZIP_PATH.unlink()
    with zipfile.ZipFile(ZIP_PATH, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in files:
            zf.write(f, f.relative_to(SUBMISSION_DIR))
    print(f"Zip: {ZIP_PATH}")
    for f in sorted(files):
        print(f"  {f.relative_to(SUBMISSION_DIR)}  ({f.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
